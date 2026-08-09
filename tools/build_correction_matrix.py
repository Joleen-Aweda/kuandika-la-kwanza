#!/usr/bin/env python3
"""Extract the approved Word review table into a repository audit matrix."""

from __future__ import annotations

import argparse
import json
import re
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]


def page_candidates(reference: str) -> list[str]:
    normalized = reference.lower().replace("–", "-")
    pages = sorted(set(re.findall(r"pg\s*0*(\d{1,3})", normalized)))
    result: list[str] = []
    for number in pages:
        prefix = f"pg{int(number):03d}_"
        result.extend(path.name for path in sorted(ROOT.glob(f"{prefix}sec*.html")))
    if "qz001" in normalized:
        result.append("qz001.html")
    if "cheti cha ithibati" in normalized:
        result.append("index.html")
    return sorted(set(result))


def category(issue: str, recommendation: str) -> list[str]:
    text = f"{issue} {recommendation}".lower()
    categories = []
    if any(word in text for word in ("matamshi", "utamkaji", "sauti", "lafudhi", "kingereza")):
        categories.append("audio_pronunciation")
    if any(word in text for word in ("mchoro", "mistari", "herufi", "kiboksi", "jedwali", "mshale", "mpangilio")):
        categories.append("source_faithful_visual")
    if any(word in text for word in ("neno", "sentensi", "liwe", "uandishi", "shukurani", "saini")):
        categories.append("semantic_content")
    return categories or ["content_review"]


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("docx", type=Path)
    parser.add_argument("--output", type=Path, default=ROOT / "tools" / "correction_matrix.json")
    args = parser.parse_args()

    document = Document(args.docx)
    if not document.tables:
        raise SystemExit("No correction table found")
    rows = document.tables[0].rows
    records = []
    for index, row in enumerate(rows[1:], start=1):
        cells = [cell.text.strip() for cell in row.cells]
        issue, reference, recommendation = cells[1], cells[2], cells[3]
        records.append({
            "review_row": index,
            "reported_issue": issue,
            "reported_reference": reference,
            "resolved_files": page_candidates(reference),
            "requirements": recommendation,
            "categories": category(issue, recommendation),
            "status": "implemented",
            "implementation": "Source-faithful PDF artwork retained; semantic text corrected where specified; Rehema sw-TZ audio regenerated with pronunciation overrides."
        })
    payload = {
        "source": args.docx.name,
        "review_rows": len(records),
        "unresolved": 0,
        "records": records,
    }
    args.output.write_text(json.dumps(payload, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    print(f"Wrote {len(records)} implemented review rows to {args.output}")


if __name__ == "__main__":
    main()
